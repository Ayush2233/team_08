import json
import docx
import fitz
import pdfplumber

def load_json(file_path: str) -> dict:
    with open(file_path, 'r') as f:
        return json.load(f)

def parse_docx(file_path: str) -> str:
    """
    Parse and return text content from a DOCX file.
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return "\n".join(full_text)

def parse_rfp_document(file_path: str) -> str:
    with open(file_path, 'r') as f:
        return f.read()

def parse_pdf(file_path: str) -> str:
    """
    Parse and return text content from a PDF file using PyMuPDF.
    """
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"
    return full_text


def load_rfp_text(file_path):
    all_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_text += text + "\n"
    return all_text


import fitz  # PyMuPDF
import docx  # python-docx
import os

def parse_pdf_streamlit(file_obj):
    try:
        file_name = file_obj.name.lower()
        
        if file_name.endswith(".pdf"):
            file_bytes = file_obj.read()
            doc = fitz.open("pdf", file_bytes)
            full_text = ""
            for page in doc:
                full_text += page.get_text() + "\n"
            return full_text
        
        elif file_name.endswith(".docx"):
            docx_file = docx.Document(file_obj)
            full_text = ""
            for para in docx_file.paragraphs:
                full_text += para.text + "\n"
            return full_text

        else:
            raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

    except Exception as e:
        raise Exception(f"Error parsing file: {e}")
