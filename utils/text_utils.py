from docx import Document
from docx.shared import Pt
from markdown2 import markdown
from bs4 import BeautifulSoup
import re

def chunk_text(text: str, chunk_size: int = 1000) -> list:

    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
    return chunks

def dynamic_chunk_text(text: str) -> list:
    
    pattern = re.compile(r'(?=\n?\d+\.\s+[A-Z])')
    
    sections = pattern.split(text)
    
    sections = [section.strip() for section in sections if section.strip()]
    return sections

def markdown_to_docx(markdown_text, output_path="generated_proposal.docx"):
    
    html = markdown(markdown_text)

    soup = BeautifulSoup(html, "html.parser")

    # Create a Word document
    doc = Document()

    for element in soup.recursiveChildGenerator():
        if element.name == "h1":
            doc.add_heading(element.text.strip(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.text.strip(), level=2)
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph("• " + li.text.strip(), style="List Bullet")
        elif element.name == "ol":
            for idx, li in enumerate(element.find_all("li"), 1):
                doc.add_paragraph(f"{idx}. {li.text.strip()}", style="List Number")
        elif element.name == "p":
            text = re.sub(r"\n", "", element.text.strip())
            if text:
                doc.add_paragraph(text)

    doc.save(output_path)
    return output_path